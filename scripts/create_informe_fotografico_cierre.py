from pathlib import Path
from datetime import date

from PIL import Image, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Informe_fotografico_cierre_perimetral.docx"
PHOTO_ROOT = Path(r"C:\Users\johnn\Downloads\Photos-1-001 (1)")
TMP = ROOT / "docs" / "_render_assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111827"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = 120


def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA, borders=True):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    borders_node = tbl_pr.find(qn("w:tblBorders"))
    if borders_node is None:
        borders_node = OxmlElement("w:tblBorders")
        tbl_pr.append(borders_node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders_node.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders_node.append(el)
        el.set(qn("w:val"), "single" if borders else "nil")
        if borders:
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "D0D5DD")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=MUTED)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("ENLACE CONSTRUCTOR  |  INFORME TÉCNICO")
    set_font(left, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    label = p.add_run("Cierre perimetral · Página ")
    set_font(label, size=8.5, color=MUTED)
    add_page_field(p)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_previous = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_font(a, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        set_font(b)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA], borders=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_font(r, size=10.5, color=BLACK)
    return table


def image_dimensions(path, max_width_in, max_height_in):
    with Image.open(path) as im:
        width, height = im.size
    ratio = min(max_width_in / width, max_height_in / height)
    return Inches(width * ratio), Inches(height * ratio)


def optimized_image(path):
    """Create a compact report copy; keeps source files untouched."""
    TMP.mkdir(parents=True, exist_ok=True)
    source = Path(path)
    target = TMP / f"{source.stem}_report.jpg"
    with Image.open(source) as im:
        rgb = ImageOps.exif_transpose(im).convert("RGB")
        rgb.thumbnail((1800, 1800), Image.Resampling.LANCZOS)
        rgb.save(target, "JPEG", quality=84, optimize=True, progressive=True)
    return target


def add_image_to_cell(cell, path, caption, max_width=3.0, max_height=2.12):
    path = optimized_image(path)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    width, height = image_dimensions(path, max_width, max_height)
    run.add_picture(str(path), width=width, height=height)
    cap = cell.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.add_run(caption)


def add_photo_grid(doc, entries):
    rows = (len(entries) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    set_table_geometry(table, [4680, 4680], borders=False)
    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(entries):
                path, caption = entries[idx]
                add_image_to_cell(cell, path, caption)
            else:
                cell.text = ""
            idx += 1
    return table


def add_scope_table(doc):
    rows = [
        ("1", "Demolición controlada de panderetas", "Por m²; incluye picado y caída controlada. No incluye traslado interno ni retiro final."),
        ("2", "Retiro de todos los postes de hormigón", "Por unidad; se retiran por fisuras, daño en bases y pérdida de confiabilidad."),
        ("3", "Demolición selectiva de dados", "Sólo los que interfieran, estén dañados o no permitan anclaje compatible."),
        ("4", "Traslado interno de escombros", "Por m³, desde el frente de trabajo al punto de acopio/carguío; recorrido en pendiente de hasta 25 m."),
        ("5", "Carguío, camión tolva y botadero", "Por viaje; retiro a disposición autorizada."),
        ("6", "Dados reutilizables o fundaciones nuevas", "Alternativas separadas: platina y anclaje químico sobre dado apto, o dado nuevo de hormigón armado."),
        ("7", "Postes y paños de cierre", "Poste mínimo 40×40×4 mm; ACMA 1G galvanizada; marco angular 25×25×3 o 4 mm."),
        ("8", "Paños laterales escalonados", "Paños cortos aproximados de 1,50 m, con mayor densidad de postes para adaptarse a la pendiente."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [720, 3000, 5640], borders=True)
    headers = ("Ítem", "Partida", "Criterio de medición / alcance")
    for idx, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=9.5, bold=True, color=INK)
    for item, name, desc in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((item, name, desc)):
            set_cell_width(cells[idx], (720, 3000, 5640)[idx])
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9.2, bold=(idx == 1))
    return table


def add_metadata(doc):
    rows = [
        ("Objeto", "Retiro y reposición de cierre perimetral dañado por temporal"),
        ("Extensión referencial", "62 m lineales: frente 33 m + dos laterales de 14,5 m"),
        ("Altura existente", "Aproximadamente 2,0 m"),
        ("Condición de acceso", "Pendiente pronunciada; traslado interno aproximado de 25 m hasta el carguío"),
        ("Fecha del informe", "23 de julio de 2026"),
        ("Carácter", "Inspección visual preliminar para presupuesto; no reemplaza cálculo ni informe estructural"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 6660], borders=True)
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=9.5, bold=True, color=INK)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.5)
    return table


def add_page_break(doc):
    doc.add_page_break()


def prepare_webp():
    TMP.mkdir(parents=True, exist_ok=True)
    source = Path(r"C:\Users\johnn\Downloads\imagen_acmanet_producto_mallafor_cerco_galvanizado_1g9_3g9_01.webp")
    target = TMP / "referencia_mallafor.png"
    with Image.open(source) as im:
        im.convert("RGB").save(target, "PNG")
    return target


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    # First page: memo masthead pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("INFORME FOTOGRÁFICO TÉCNICO")
    set_font(r, size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Retiro y reposición de cierre perimetral")
    set_font(r, size=14, bold=True, color=DARK_BLUE)

    add_metadata(doc)

    add_heading(doc, "Conclusión ejecutiva", 1)
    add_callout(
        doc,
        "RECOMENDACIÓN CRÍTICA",
        "Retirar la totalidad de los postes prefabricados de hormigón. Las fotografías muestran fisuras, "
        "desprendimientos, daño en las bases, desplazamientos y pérdida de confinamiento; por inspección visual "
        "no existe fundamento suficiente para reutilizarlos como soporte del nuevo cierre.",
        fill="FDECEC",
        accent=RED,
    )
    add_body(
        doc,
        "Los dados de fundación se tratarán de forma distinta: podrán conservarse cuando no interfieran y, "
        "para una eventual reutilización, deberán encontrarse íntegros, estables y permitir distancias de borde, "
        "profundidad y anclajes compatibles con la nueva platina. Los dados dañados o incompatibles se demolerán "
        "selectivamente y se reemplazarán.",
    )

    overview = optimized_image(PHOTO_ROOT / "IMG_20260721_165722701.jpg")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    width, height = image_dimensions(overview, 6.45, 2.55)
    p.add_run().add_picture(str(overview), width=width, height=height)
    cap = doc.add_paragraph(style="Caption")
    cap.add_run("Fotografía 1. Vista general del cierre, la pendiente y los tramos afectados.")

    add_page_break(doc)

    add_heading(doc, "1. Condición observada y riesgos de ejecución", 1)
    add_body(
        doc,
        "El temporal provocó el colapso de paños de pandereta y dejó otros tramos parcialmente en pie. "
        "El terreno presenta pendiente fuerte y material distribuido aguas abajo, por lo que la demolición "
        "deberá realizarse por sectores cortos, con caída controlada, despeje inmediato y separación del frente "
        "de montaje.",
    )
    add_photo_grid(
        doc,
        [
            (
                PHOTO_ROOT / "IMG_20260721_171123786.jpg",
                "Fotografía 2. Paños colapsados sobre la pendiente y recorrido manual de extracción.",
            ),
            (
                PHOTO_ROOT / "IMG_20260721_170150796.jpg",
                "Fotografía 3. Elementos desplazados y postes fuera de posición.",
            ),
            (
                PHOTO_ROOT / "IMG_20260721_170011158.jpg",
                "Fotografía 4. Tramo aún en pie junto a talud; requiere demolición controlada.",
            ),
            (
                PHOTO_ROOT / "IMG_20260721_170022038.jpg",
                "Fotografía 5. Acumulación de placas y restricciones de tránsito en el frente.",
            ),
        ],
    )
    add_heading(doc, "Criterios operacionales", 2)
    add_bullet(doc, "Trabajar con cuadrillas diferenciadas de demolición/logística y fabricación-montaje.")
    add_bullet(doc, "Demoler y reponer por tramos cortos para reducir el tiempo de apertura del perímetro.")
    add_bullet(doc, "Medir el traslado interno desde el frente hasta el punto de acopio/carguío, sin condicionar el método a un equipo específico.")
    add_bullet(doc, "Mantener control local del área, coordinación con residentes y protección de terceros durante toda la faena.")

    add_page_break(doc)

    add_heading(doc, "2. Daño crítico en postes existentes", 1)
    add_body(
        doc,
        "Las siguientes evidencias corresponden a zonas de apoyo y encuentros entre postes y placas. Se observan "
        "fisuras longitudinales y transversales, pérdida de material, desconfinamiento y apoyos alterados. "
        "Aunque una fotografía no permite cuantificar la capacidad resistente remanente, sí permite descartar "
        "su aceptación automática como soporte de un sistema nuevo.",
    )
    add_photo_grid(
        doc,
        [
            (
                PHOTO_ROOT / "IMG_20260721_165656471.jpg",
                "Fotografía 6. Fisuración y desprendimiento en la base de un poste.",
            ),
            (
                PHOTO_ROOT / "IMG_20260721_165959976.jpg",
                "Fotografía 7. Fisura longitudinal abierta en zona basal.",
            ),
            (
                PHOTO_ROOT / "IMG_20260721_165913840.jpg",
                "Fotografía 8. Fisura transversal y degradación en el encuentro.",
            ),
            (
                PHOTO_ROOT / "IMG_20260721_165317383.jpg",
                "Fotografía 9. Pérdida de sección y exposición del apoyo inferior.",
            ),
        ],
    )
    add_callout(
        doc,
        "Decisión para el presupuesto",
        "La partida considera retirar todos los postes de hormigón. La demolición de dados queda separada y se "
        "medirá sólo cuando el dado interfiera, esté dañado o no sea compatible con el nuevo anclaje.",
        fill=LIGHT,
        accent=DARK_BLUE,
    )

    add_page_break(doc)

    add_heading(doc, "3. Solución de reposición propuesta", 1)
    add_body(
        doc,
        "El nuevo cierre se propone con malla electrosoldada galvanizada tipo ACMA 1G, contenida en un marco "
        "soldado de ángulo 25×25×3 mm o 25×25×4 mm. Los postes serán perfiles tubulares de acero de sección "
        "mínima 40×40×4 mm, con protección anticorrosiva en cortes y soldaduras.",
    )
    ref_png = prepare_webp()
    add_photo_grid(
        doc,
        [
            (
                Path(r"C:\Users\johnn\Downloads\R.jpg"),
                "Referencia A. Apariencia esperada del cierre metálico terminado.",
            ),
            (
                ref_png,
                "Referencia B. Malla galvanizada enmarcada y fijada a postes metálicos.",
            ),
        ],
    )
    add_heading(doc, "Modulación y fundaciones", 2)
    add_bullet(doc, "Frente: paños estándar cercanos a 3,00 m, ajustados al levantamiento definitivo.")
    add_bullet(doc, "Laterales: paños cortos cercanos a 1,50 m y mayor cantidad de postes para absorber el escalonamiento.")
    add_bullet(doc, "Dados aptos: platina soldada al poste y pernos con anclaje químico, previa verificación del sustrato y las distancias de borde.")
    add_bullet(doc, "Dados incompatibles: demolición puntual y ejecución de una nueva fundación de hormigón armado.")
    add_callout(
        doc,
        "Verificación requerida antes de fabricar",
        "Confirmar en terreno niveles, largo real de cada tramo, ubicación de redes, geometría de los dados y "
        "cantidad de paños. La sección 40×40×4 mm es el mínimo solicitado y deberá validarse según altura, "
        "separación entre postes, solicitación de viento y exposición costera.",
        fill="FFF7E0",
        accent=GOLD,
    )

    add_page_break(doc)

    add_heading(doc, "4. Partidas y APU preparados para el presupuesto", 1)
    add_body(
        doc,
        "Las partidas se incorporaron separando actividades que tienen rendimientos y riesgos distintos. "
        "Esta estructura evita esconder la logística de la pendiente dentro de la demolición y permite ajustar "
        "cantidades sin duplicar costos.",
    )
    add_scope_table(doc)
    add_heading(doc, "Secuencia recomendada", 2)
    add_bullet(doc, "Replantear y confirmar la modulación final de postes y paños.")
    add_bullet(doc, "Desmontar o demoler un tramo corto, retirar poste y despejar el material.")
    add_bullet(doc, "Evaluar cada dado expuesto: conservar, reutilizar con anclaje o demoler selectivamente.")
    add_bullet(doc, "Instalar poste y paño nuevo antes de ampliar el siguiente frente.")
    add_bullet(doc, "Trasladar escombros al punto de carguío y coordinar su retiro en camión tolva.")

    add_heading(doc, "Alcances y limitaciones", 2)
    add_body(
        doc,
        "No se incorpora instalación de faena independiente ni cierre provisorio continuo, debido a que el "
        "condominio está en uso y la reposición se ejecutará progresivamente. Esto no elimina las obligaciones "
        "de señalización, control del frente, elementos de protección personal y medidas para impedir la caída "
        "de materiales o el ingreso de terceros.",
    )
    add_body(
        doc,
        "Este documento registra una inspección visual y apoya la elaboración de la oferta. Las dimensiones, "
        "cantidades, condición de fundaciones, estabilidad del talud y solución definitiva de anclaje deben "
        "verificarse en terreno antes de contratar materiales o iniciar fabricación.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
